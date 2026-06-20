import json
import base64
import io
import re
import os
from datetime import datetime
from http.server import BaseHTTPRequestHandler
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Design tokens (fra Blåfall design handoff)
C_DARK        = RGBColor(0x1f, 0x48, 0x56)   # mørkeste teal
C_TEAL        = RGBColor(0x2f, 0x5f, 0x6f)   # teal (stedsnavn)
C_MID         = RGBColor(0x3d, 0x6e, 0x7f)   # primær teal (overskrifter, strek)
C_MUTED       = RGBColor(0x5b, 0x7d, 0x8a)   # dempet teal (eyebrow, dato)
C_GREY        = RGBColor(0x8a, 0xa3, 0xac)   # lys grå (header/footer tekst)
C_PLACEHOLDER = RGBColor(0xC0, 0x39, 0x2B)

LOGO_PATH = os.path.join(os.path.dirname(__file__), '..', 'assets', 'blaafall_logo.png')

NORSK_MAANEDER = [
    'januar', 'februar', 'mars', 'april', 'mai', 'juni',
    'juli', 'august', 'september', 'oktober', 'november', 'desember'
]


def _pBdr(paragraph, side, color_hex, size, space='4'):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = OxmlElement('w:pBdr')
        pPr.append(pBdr)
    el = OxmlElement(f'w:{side}')
    el.set(qn('w:val'), 'single')
    el.set(qn('w:sz'), str(size))
    el.set(qn('w:space'), space)
    el.set(qn('w:color'), color_hex)
    pBdr.append(el)


def _no_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for border in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        el = OxmlElement(f'w:{border}')
        el.set(qn('w:val'), 'none')
        tblBorders.append(el)
    tblPr.append(tblBorders)
    # Fjern også borders på alle celler
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                el = OxmlElement(f'w:{border}')
                el.set(qn('w:val'), 'none')
                tcBorders.append(el)
            tcPr.append(tcBorders)


def extract_metadata(text):
    anlegg = 'Kraftverk'
    kommune = ''
    fylke = ''

    # Prosjektnavn fra første bold-linje
    bold_matches = re.findall(r'\*\*([^*]+)\*\*', text[:1000])
    if bold_matches:
        first = bold_matches[0]
        m = re.search(r'bygging av (.+)', first, re.IGNORECASE)
        anlegg = m.group(1).strip() if m else first

    # Kommune og fylke fra SAMMENDRAG-tabell
    m = re.search(r'\|\s*Kommune\s*\|\s*([^|]+)\s*\|', text, re.IGNORECASE)
    if m:
        kommune = m.group(1).strip()
    m = re.search(r'\|\s*Fylke\s*\|\s*([^|]+)\s*\|', text, re.IGNORECASE)
    if m:
        fylke = m.group(1).strip()

    # Dato fra bold månedsnavn i teksten
    pattern = r'\*\*(' + '|'.join(NORSK_MAANEDER) + r')\s+(\d{4})\*\*'
    m = re.search(pattern, text, re.IGNORECASE)
    if m:
        maaned = m.group(1).capitalize()
        aar = m.group(2)
    else:
        now = datetime.now()
        maaned = NORSK_MAANEDER[now.month - 1].capitalize()
        aar = str(now.year)

    return anlegg, kommune, fylke, maaned, aar


def _setup_header(section, anlegg):
    header = section.header
    header.is_linked_to_previous = False

    for p in header.paragraphs:
        p.clear()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    tbl = header.add_table(rows=1, cols=2, width=Cm(16.5))
    _no_borders(tbl)
    tbl.cell(0, 0).width = Cm(7)
    tbl.cell(0, 1).width = Cm(9.5)

    # Logo venstre
    logo_p = tbl.cell(0, 0).paragraphs[0]
    logo_p.paragraph_format.space_before = Pt(0)
    logo_p.paragraph_format.space_after = Pt(0)
    if os.path.exists(LOGO_PATH):
        logo_p.add_run().add_picture(LOGO_PATH, width=Cm(4.5))
    else:
        run = logo_p.add_run('BLÅFALL')
        run.bold = True
        run.font.color.rgb = C_DARK
        run.font.size = Pt(14)

    # Prosjektnavn høyre
    title_p = tbl.cell(0, 1).paragraphs[0]
    title_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    title_p.paragraph_format.space_before = Pt(10)
    title_p.paragraph_format.space_after = Pt(0)
    run = title_p.add_run(anlegg.upper())
    run.font.name = 'Calibri'
    run.font.size = Pt(8)
    run.font.bold = True
    run.font.color.rgb = C_GREY

    # Horisontal strek
    rule = header.add_paragraph()
    rule.paragraph_format.space_before = Pt(4)
    rule.paragraph_format.space_after = Pt(0)
    _pBdr(rule, 'bottom', '3d6e7f', 16)


def _setup_footer(section, dato_str):
    footer = section.footer
    footer.is_linked_to_previous = False

    for p in footer.paragraphs:
        p.clear()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    # Tynn strek over footer
    rule = footer.paragraphs[0]
    rule.paragraph_format.space_before = Pt(0)
    rule.paragraph_format.space_after = Pt(6)
    _pBdr(rule, 'top', 'cdd9de', 4)

    tbl = footer.add_table(rows=1, cols=2, width=Cm(16.5))
    _no_borders(tbl)
    tbl.cell(0, 0).width = Cm(8)
    tbl.cell(0, 1).width = Cm(8.5)

    left_p = tbl.cell(0, 0).paragraphs[0]
    left_p.paragraph_format.space_before = Pt(0)
    left_p.paragraph_format.space_after = Pt(0)
    run = left_p.add_run('BLÅFALL AS')
    run.font.name = 'Calibri'
    run.font.size = Pt(8)
    run.font.bold = True
    run.font.color.rgb = C_GREY

    right_p = tbl.cell(0, 1).paragraphs[0]
    right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right_p.paragraph_format.space_before = Pt(0)
    right_p.paragraph_format.space_after = Pt(0)
    run = right_p.add_run(f'SØKNAD OM KONSESJON · {dato_str.upper()}')
    run.font.name = 'Calibri'
    run.font.size = Pt(8)
    run.font.bold = True
    run.font.color.rgb = C_GREY


def _build_cover(doc, anlegg, kommune, fylke, maaned, aar):
    def cp(text, size, color, bold=False, italic=False, space_before=0, space_after=0):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        run.font.color.rgb = color
        return p

    cp('BLÅFALL AS', 13, C_MUTED, bold=True, space_before=44)
    cp(anlegg, 34, C_DARK, bold=True, space_before=14)
    cp('— i —', 12, C_MUTED, italic=True, space_before=16)

    if kommune or fylke:
        location = f'{kommune} · {fylke}' if (kommune and fylke) else (kommune or fylke)
        cp(location, 15, C_TEAL, bold=True, space_before=14)

    # Bildeplass
    img_p = cp('[Bilde av anlegget]', 10, C_GREY, italic=True, space_before=30, space_after=0)
    img_p.paragraph_format.space_after = Pt(30)
    for side in ['top', 'bottom', 'left', 'right']:
        _pBdr(img_p, side, 'cdd9de', 4)

    cp('Søknad om konsesjon', 22, C_DARK, bold=True, space_before=24)
    cp(f'{maaned.upper()} {aar}', 12, C_MUTED, bold=True, space_before=10)

    doc.add_page_break()


def add_formatted_runs(paragraph, text):
    parts = re.split(r'(\*\*[^*]+?\*\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(part)
        if '[FYLL INN' in part:
            run.font.color.rgb = C_PLACEHOLDER
            run.bold = True


def render_content(doc, text):
    table_rows = []
    in_table = False

    def flush_table():
        if not table_rows:
            return
        max_cols = max(len(r) for r in table_rows)
        tbl = doc.add_table(rows=len(table_rows), cols=max_cols)
        tbl.style = 'Table Grid'
        for r_i, row in enumerate(table_rows):
            for c_i in range(max_cols):
                cell_text = row[c_i] if c_i < len(row) else ''
                cell = tbl.cell(r_i, c_i)
                cell.text = cell_text
                if r_i == 0:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.bold = True
        table_rows.clear()

    for line in text.split('\n'):
        s = line.strip()

        if s.startswith('|'):
            cells = [c.strip() for c in s.split('|')[1:-1]]
            if not all(set(c) <= set('-: ') for c in cells):
                in_table = True
                table_rows.append(cells)
            continue

        if in_table:
            flush_table()
            in_table = False

        if not s:
            continue

        if s == '---':
            doc.add_page_break()
        elif s.startswith('# '):
            doc.add_heading(s[2:], level=1)
        elif s.startswith('## '):
            doc.add_heading(s[3:], level=2)
        elif s.startswith('### '):
            doc.add_heading(s[4:], level=3)
        elif s.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_runs(p, s[2:])
        elif re.match(r'^\d+\.\s', s):
            p = doc.add_paragraph(style='List Number')
            add_formatted_runs(p, re.sub(r'^\d+\.\s*', '', s))
        else:
            p = doc.add_paragraph()
            add_formatted_runs(p, s)

    if in_table:
        flush_table()


def build_docx(text):
    doc = Document()

    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)
    for level, size in [(1, 14), (2, 13), (3, 12)]:
        style = doc.styles[f'Heading {level}']
        style.font.name = 'Calibri'
        style.font.size = Pt(size)
        style.font.color.rgb = C_MID
        style.font.bold = True

    anlegg, kommune, fylke, maaned, aar = extract_metadata(text)
    dato_str = f'{maaned} {aar}'

    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(4.2)
    section.bottom_margin = Cm(3.0)
    section.left_margin = Cm(3.5)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)

    _setup_header(section, anlegg)
    _setup_footer(section, dato_str)
    _build_cover(doc, anlegg, kommune, fylke, maaned, aar)

    # Innhold: start fra NVE-adressen (etter tittelblokken)
    nve_match = re.search(r'NVE\s*[–-]', text)
    if nve_match:
        content_text = text[nve_match.start():]
    else:
        parts = text.split('---', 1)
        content_text = parts[1] if len(parts) > 1 else text

    render_content(doc, content_text)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


class handler(BaseHTTPRequestHandler):
    def do_POST(self):
        try:
            length = int(self.headers.get("Content-Length", 0))
            body = json.loads(self.rfile.read(length))

            docx_bytes = build_docx(body.get("document", ""))
            docx_b64 = base64.b64encode(docx_bytes).decode('utf-8')

            self.send_response(200)
            self.send_header("Content-Type", "application/json")
            self.send_header("Access-Control-Allow-Origin", "*")
            self.end_headers()
            self.wfile.write(json.dumps({"docx": docx_b64}).encode())

        except Exception as e:
            self.send_response(500)
            self.send_header("Content-Type", "application/json")
            self.send_header("Access-Control-Allow-Origin", "*")
            self.end_headers()
            self.wfile.write(json.dumps({"error": str(e)}).encode())

    def do_OPTIONS(self):
        self.send_response(200)
        self.send_header("Access-Control-Allow-Origin", "*")
        self.send_header("Access-Control-Allow-Methods", "POST, OPTIONS")
        self.send_header("Access-Control-Allow-Headers", "Content-Type")
        self.end_headers()
